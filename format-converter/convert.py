#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用格式转换工具 v2.0
支持 Markdown、Word、PDF 等多种格式的互相转换
优化版：不依赖外部工具，支持中文
"""

import os
import sys
import argparse
import re
from pathlib import Path

def get_file_extension(filepath):
    """获取文件扩展名"""
    return Path(filepath).suffix.lower()

def convert_file(input_path, output_path=None, output_format=None):
    """
    通用文件转换函数
    """
    input_ext = get_file_extension(input_path)

    if output_format:
        output_ext = '.' + output_format.lstrip('.')
    elif output_path:
        output_ext = get_file_extension(output_path)
    else:
        raise ValueError("Must specify output path or format")

    if not output_path:
        input_dir = os.path.dirname(input_path)
        input_basename = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(input_dir, f"{input_basename}{output_ext}")

    conversion_key = f"{input_ext[1:]}_to_{output_ext[1:]}"

    converters = {
        'md_to_docx': md_to_docx,
        'md_to_pdf': md_to_pdf_v2,
        'docx_to_md': docx_to_md_v2,
        'docx_to_pdf': docx_to_pdf_v2,
        'pdf_to_md': pdf_to_md,
        'pdf_to_docx': pdf_to_docx,
        'txt_to_md': txt_to_md,
        'txt_to_docx': txt_to_docx,
    }

    converter = converters.get(conversion_key)
    if not converter:
        raise ValueError(f"Unsupported conversion: {input_ext} -> {output_ext}")

    print(f"[INFO] Converting {input_ext} -> {output_ext}...")
    result = converter(input_path, output_path)
    print(f"[OK] {result}")
    return result

# ==================== 辅助函数 ====================

def register_chinese_font():
    """注册中文字体到 reportlab"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Windows 中文字体路径
        font_paths = [
            ('C:/Windows/Fonts/msyh.ttc', 'Microsoft-YaHei'),
            ('C:/Windows/Fonts/msyhbd.ttc', 'Microsoft-YaHei-Bold'),
            ('C:/Windows/Fonts/simhei.ttf', 'SimHei'),
            ('C:/Windows/Fonts/simsun.ttc', 'SimSun'),
        ]

        registered = False
        for font_path, font_name in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    registered = True
                except:
                    pass

        return registered
    except ImportError:
        return False
    except Exception:
        return False

# ==================== Markdown → Word ====================

def md_to_docx(md_path, docx_path):
    """Markdown 转 Word"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        if line.startswith('# '):
            title = doc.add_heading(line[2:].strip(), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif '**' in line:
            p = doc.add_paragraph()
            last_end = 0
            for match in re.finditer(r'\*\*(.*?)\*\*', line):
                if match.start() > last_end:
                    p.add_run(line[last_end:match.start()])
                p.add_run(match.group(1)).bold = True
                last_end = match.end()
            if last_end < len(line):
                p.add_run(line[last_end:])
        elif line.startswith(('- ', '* ')):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    return docx_path

# ==================== Markdown → PDF (优化版) ====================

def md_to_pdf_v2(md_path, pdf_path):
    """Markdown 转 PDF（不依赖外部工具）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.pdfbase import pdfmetrics

        # 读取 Markdown
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 创建 PDF
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # 注册中文字体
        has_chinese = register_chinese_font()

        # 创建样式
        styles = getSampleStyleSheet()

        if has_chinese:
            styles.add(ParagraphStyle(
                name='ChineseNormal',
                fontName='Microsoft-YaHei',
                fontSize=11,
                leading=16,
                alignment=TA_LEFT,
            ))
            styles.add(ParagraphStyle(
                name='ChineseHeading1',
                fontName='Microsoft-YaHei',
                fontSize=18,
                leading=24,
                alignment=TA_CENTER,
                spaceAfter=12,
            ))
            styles.add(ParagraphStyle(
                name='ChineseHeading2',
                fontName='Microsoft-YaHei',
                fontSize=14,
                leading=20,
                spaceAfter=10,
            ))
            normal_style = styles['ChineseNormal']
            heading1_style = styles['ChineseHeading1']
            heading2_style = styles['ChineseHeading2']
        else:
            normal_style = styles['Normal']
            heading1_style = styles['Heading1']
            heading2_style = styles['Heading2']

        # 构建内容
        story = []
        lines = md_content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.3*cm))
                continue

            # 处理标题
            if line.startswith('# '):
                text = line[2:].strip()
                story.append(Paragraph(text, heading1_style))
            elif line.startswith('## '):
                text = line[3:].strip()
                story.append(Paragraph(text, heading2_style))
            elif line.startswith('### '):
                text = line[4:].strip()
                story.append(Paragraph(text, heading2_style))
            else:
                # 处理粗体
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                story.append(Paragraph(line, normal_style))

        doc.build(story)
        return pdf_path

    except ImportError:
        print("[WARN] reportlab not installed, falling back to pandoc method...")
        return md_to_pdf_pandoc(md_path, pdf_path)
    except Exception as e:
        print(f"[WARN] Direct conversion failed: {e}")
        print("[INFO] Trying alternative method...")
        return md_to_pdf_pandoc(md_path, pdf_path)

def md_to_pdf_pandoc(md_path, pdf_path):
    """使用 pandoc 转换（备用方法）"""
    try:
        import pypandoc
        pypandoc.convert_file(md_path, 'pdf', outputfile=pdf_path)
        return pdf_path
    except ImportError:
        raise Exception("PDF conversion requires reportlab. Install: pip install reportlab")

# ==================== Word → Markdown (优化版) ====================

def docx_to_md_v2(docx_path, md_path):
    """Word 转 Markdown（优化版，支持中文标题）"""
    from docx import Document

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        # 检测标题（支持中英文样式名）
        style_name = para.style.name.lower()

        if 'heading 1' in style_name or '标题 1' in style_name or 'title' in style_name:
            lines.append(f'## {text}')
        elif 'heading 2' in style_name or '标题 2' in style_name:
            lines.append(f'### {text}')
        elif 'heading 3' in style_name or '标题 3' in style_name:
            lines.append(f'#### {text}')
        else:
            lines.append(text)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    return md_path

# ==================== Word → PDF (优化版) ====================

def docx_to_pdf_v2(docx_path, pdf_path):
    """Word 转 PDF（优化版：先转 md 再转 pdf，不依赖外部工具）"""
    try:
        # 方法1：先转 Markdown，再转 PDF
        temp_md = pdf_path.replace('.pdf', '_temp.md')
        docx_to_md_v2(docx_path, temp_md)
        result = md_to_pdf_v2(temp_md, pdf_path)

        # 删除临时文件
        if os.path.exists(temp_md):
            os.remove(temp_md)

        return result
    except Exception as e:
        # 方法2：尝试使用 pandoc
        try:
            import pypandoc
            pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
            return pdf_path
        except ImportError:
            raise Exception(f"Word to PDF failed: {e}\nInstall: pip install reportlab")

# ==================== PDF → Markdown ====================

def pdf_to_md(pdf_path, md_path):
    """PDF 转 Markdown"""
    try:
        import pdfplumber

        with pdfplumber.open(pdf_path) as pdf:
            text_content = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(text_content))

        return md_path
    except ImportError:
        raise Exception("PDF to Markdown requires pdfplumber. Install: pip install pdfplumber")

# ==================== PDF → Word ====================

def pdf_to_docx(pdf_path, docx_path):
    """PDF 转 Word"""
    try:
        from pdf2docx import Converter

        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()

        return docx_path
    except ImportError:
        raise Exception("PDF to Word requires pdf2docx. Install: pip install pdf2docx")

# ==================== TXT → Markdown/Word ====================

def txt_to_md(txt_path, md_path):
    """纯文本转 Markdown"""
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return md_path

def txt_to_docx(txt_path, docx_path):
    """纯文本转 Word"""
    from docx import Document

    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(docx_path)
    return docx_path

# ==================== 命令行接口 ====================

def main():
    parser = argparse.ArgumentParser(
        description='Universal format converter v2.0 - No external tools required',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  %(prog)s input.md output.docx          # Markdown -> Word
  %(prog)s input.docx output.pdf          # Word -> PDF (NEW: no external tools!)
  %(prog)s input.pdf output.md            # PDF -> Markdown
  %(prog)s --from md --to pdf input.md    # Specify formats
        '''
    )

    parser.add_argument('input', help='Input file')
    parser.add_argument('output', nargs='?', help='Output file')
    parser.add_argument('--to', dest='output_format', help='Output format (md, docx, pdf)')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"[ERROR] Input file not found: {args.input}")
        sys.exit(1)

    try:
        result = convert_file(args.input, args.output, args.output_format)
        print(f"[SUCCESS] {result}")
    except Exception as e:
        print(f"[ERROR] {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()
